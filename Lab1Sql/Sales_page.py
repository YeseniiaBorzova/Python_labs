from Sales import Sales
from page import Page
from tkinter import *
from functools import partial
from docx import Document
import openpyxl
from tkinter import messagebox
from datetime import datetime
import re


class SalesPage(Page):

    def __init__(self, *args, **kwargs):
        Page.__init__(self, *args, **kwargs)

        window = Toplevel(self)
        window.title("Sales window")
        window.state('zoomed')

        all_sales = Frame(window)
        all_sales .pack(side=LEFT, fill=Y)
        Label(all_sales , text="All avaliable sellers(double click to delete)").pack()
        self.allsales_listbox = Listbox(all_sales , width=80)
        self.allsales_listbox.bind('<Double-1>', self.all_list_on_click)
        self.sales_list = Sales.show_all_sales().split("\n")
        self.sales_list.pop(-1)
        for sale in self.sales_list:
            self.allsales_listbox.insert(END, sale)
        self.allsales_listbox.pack(side="top", fill="both")
        all_docx = Button(all_sales, text="Export info in docx about all sales", command=self.form_docx_report())
        all_docx.pack()
        all_exel = Button(all_sales, text="Export info in exel about all sales", command = self.form_exel_report())
        all_exel.pack()

        search_sale = Frame(window)
        search_sale.pack(side=LEFT, fill=Y)
        Label(search_sale, text="Search sale by date").pack()
        date_entry = Entry(search_sale)
        date_entry.pack()
        self.find_listbox = Listbox(search_sale, width=80)
        self.find_listbox.bind('<Double-1>', self.find_list_on_click)
        find_sale = partial(self.search_sale_on_click, self.find_listbox, date_entry)
        search_btn = Button(search_sale, text="Search", command=find_sale)
        search_btn.pack()
        self.find_listbox.pack()

        add_sale = Frame(window)
        add_sale.pack(side=LEFT, fill=Y)
        Label(add_sale, text="Add new sale").grid(row=0, column=0, columnspan=3)
        seller_id = Label(add_sale, text="Enter sellers` id:")
        prod_label = Label(add_sale, text="Enter products splited by coma:")
        amount_label = Label(add_sale, text="Enter amounts of each product by coma:")
        id_entry = Entry(add_sale)
        prod_entry = Entry(add_sale)
        amount_entry = Entry(add_sale)
        new_sale_add = partial(self.add_sale, id_entry, prod_entry, amount_entry, self.allsales_listbox)
        add_btn = Button(add_sale, text="Add", command=new_sale_add)
        seller_id.grid(row=1, column=0)
        prod_label.grid(row=2, column=0)
        amount_label.grid(row=3, column=0)
        id_entry.grid(row=1, column=1)
        prod_entry.grid(row=2, column=1)
        amount_entry.grid(row=3, column=1)
        add_btn.grid(row=5, column=1)

    def all_list_on_click(self, event):
        cs = self.allsales_listbox.curselection()
        self.manipulate_sales(self.sales_list[cs[0]])

    def find_list_on_click(self, event):
        cs = self.find_listbox.curselection()
        self.manipulate_sale(self.find_list)

    def update_list(self, listbox):
        listbox.delete(0, END)
        self.sales_list = Sales.show_all_sales().split('\n')
        for sale in self.sales_list:
            self.allsales_listbox.insert(END, sale)
        return

    def search_sale_on_click(self, listbox, entry_date):
        listbox.delete(0, END)
        try:
            self.find_list = Sales.find_sale_by_date(entry_date.get())
            listbox.insert(END, self.find_list)
        except IndexError:
            messagebox.showerror("No such sale error!", "Incorrect sale date or such sale doesnt exist!")
        return

    def add_sale(self, id_entry, prod_entry, amount_entry, listbox):
        prod_arr = prod_entry.get().split(',')
        amount_arr = list(map(float, amount_entry.get().split(",")))
        new_sale = Sales(int(id_entry.get()), prod_arr, amount_arr)
        new_sale.init_sale()
        id_entry.delete(0, END)
        prod_entry.delete(0, END)
        amount_entry.delete(0, END)
        self.update_list(listbox)

    def manipulate_sales(self, sale):
        word_list = re.sub("[^\w]", " ", sale).split()
        sale_id = int(word_list[1])
        seller_id = int(word_list[4])
        MsgBox = messagebox.askquestion("Delete sale", "Are you sure that you want to delete sale?", icon='warning')
        if MsgBox == 'yes':
            Sales.delete_sale_by_id(sale_id)
            self.update_list(self.allsales_listbox)
        else:
            self.update_list(self.allsales_listbox)

    def form_exel_report(self):
        wb = openpyxl.load_workbook('Reports\\All_info.xlsx')
        if 'Sales' not in wb.sheetnames:
            wb.create_sheet("Sales")
        ws = wb.get_sheet_by_name("Sales")
        ws.delete_cols(1, 5)
        ws.delete_rows(1, 100)
        for i in range(len(self.sales_list)):
            result_id = re.search('Id: (.*), seller id:', self.sales_list[i]).group(1)
            result_seller_id = re.search('seller id: (.*), ', self.sales_list[i]).group(1)
            result_date = re.search('datetime:(.*)', self.sales_list[i]).group(1)
            ws.cell(row=i + 1, column=1).value = result_id
            ws.cell(row=i + 1, column=2).value = result_seller_id
            ws.cell(row=i + 1, column=3).value = result_date


        wb.save('Reports\\All_info.xlsx')

    def form_docx_report(self):
        for i in self.sales_list:
            self.sale_export(i)

    def sale_export(self, sale):
        result_id = re.search('Id: (.*), seller id:', sale).group(1)
        result_seller_id = re.search('seller id: (.*), ', sale).group(1)
        result_date = re.search('datetime:(.*)', sale).group(1)
        document = Document()
        document.add_heading((result_id + " (Sale)"), 0)
        document.add_heading("Overall information", level=1)
        document.add_paragraph("Sale id: " + result_id)
        document.add_paragraph("Seller id: " + result_seller_id)
        document.add_paragraph("Date and time: " + result_date)

        document.save("Reports\\" +"id"+ result_id+".docx")