from PyQt5 import QtWidgets, uic
import sys
import re
from docx import Document
import openpyxl
from PyQt5.QtWidgets import QMessageBox
from PyQt5.QtGui import QStandardItem, QStandardItemModel
from PyQt5.QtWidgets import QAbstractItemView, QMainWindow, QApplication, QPushButton, QVBoxLayout, QLabel, QWidget

from Seller import Seller


class Manipulate_Seller(QMainWindow):
    def __init__(self, surname):
        super().__init__()
        uic.loadUi("manipulate_seller.ui", self)
        self.surname = surname

        self.changePerBtn.clicked.connect(self.change_percent)
        self.delSeller.clicked.connect(self.delete_seller)

    def delete_seller(self):
        Seller.delete_seller_by_surname(self.surname)
        self.close()

    def change_percent(self):
        try:
            Seller.update_percent(self.surname, float(self.enterPer.text()))
            self.enterPer.setText("")
        except:
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Critical)
            msg.setText("Error")
            msg.setInformativeText('Please insert data!')
            msg.setWindowTitle("Error")
            msg.exec_()


class Seller_Window(QMainWindow):
    def __init__(self):
        super().__init__()
        uic.loadUi("seller_form.ui", self)
        self.serchSellerBtn.clicked.connect(self.find_seller)
        self.addSeller.clicked.connect(self.add_seller)
        self.manipulate_sellerBtn.clicked.connect(self.manipulate_seller)
        self.updSellers.clicked.connect(self.update_all_sellers)
        self.sellerDocx.clicked.connect(self.form_docx_report)
        self.sellerExcel.clicked.connect(self.form_exel_report)
        self.update_all_sellers()

    def update_all_sellers(self):
        self.seller_list.setEditTriggers(QAbstractItemView.NoEditTriggers)
        model = QStandardItemModel(self.seller_list)
        for i in Seller.show_all_sellers().split("\n"):
            item = QStandardItem(i)
            model.appendRow(item)
        self.seller_list.setModel(model)

    def find_seller(self):
        try:
            model = QStandardItemModel(self.searchSeller)
            item = QStandardItem(Seller.find_seller_by_surname(self.entersur.text()))
            model.appendRow(item)
            self.searchSeller.setModel(model)
            self.entersur.setText("")
        except:
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Critical)
            msg.setText("Error")
            msg.setInformativeText('No such seller')
            msg.setWindowTitle("Error")
            msg.exec_()

    def add_seller(self):
        try:
            new_seller = Seller(self.name.text(), self.surname.text(), self.midName.text(), float(self.percent.text()))
            new_seller.insert_seller()
            self.update_all_sellers()
            self.name.setText("")
            self.surname.setText("")
            self.midName.setText("")
            self.percent.setText("")
        except:
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Critical)
            msg.setText("Error")
            msg.setInformativeText('Please insert data!')
            msg.setWindowTitle("Error")
            msg.exec_()

    def manipulate_seller(self):
        surname = self.manipulate_surname.text()
        self.window = Manipulate_Seller(surname)
        self.window.show()
        self.update_all_sellers()
        self.manipulate_surname.setText("")

    def form_exel_report(self):
        wb = openpyxl.load_workbook('Reports\\All_info.xlsx')
        if 'Sellers' not in wb.sheetnames:
            wb.create_sheet("Sellers")
        ws = wb["Sellers"]
        ws.delete_cols(1, 3)
        ws.delete_rows(1, 100)
        seller_list = Seller.show_all_sellers().split("\n")
        del seller_list[-1]
        for i in range(len(seller_list)):
            result_id = re.search('Id: (.*), Name', seller_list[i]).group(1)
            result_name = re.search('Name: (.*), comission ', seller_list[i]).group(1)
            result_percent = re.search(' percent: (.*) ', seller_list[i]).group(1)
            ws.cell(row=i + 1, column=1).value = result_id
            ws.cell(row=i + 1, column=2).value = result_name
            ws.cell(row=i + 1, column=3).value = result_percent

        wb.save('Reports\\All_info.xlsx')

    def form_docx_report(self):
        seller_list = Seller.show_all_sellers().split("\n")
        del seller_list[-1]
        for i in seller_list:
            self.seller_export(i)

    def seller_export(self, seller):
        result_id = re.search('Id: (.*), Name', seller).group(1)
        result_name = re.search('Name: (.*), comission ', seller).group(1)
        result_percent = re.search(' percent: (.*) ', seller).group(1)
        result = Seller.find_seller_sales(result_id).split("\n")
        result.pop(-1)
        document = Document()
        document.add_heading((result_id + ' ' + result_name + ' ' + "(Seller)"), 0)
        document.add_heading("Overall information", level=1)
        document.add_paragraph("Id: " + result_id)
        document.add_paragraph("Name: " + result_name)
        document.add_paragraph("Commission percent: " + result_percent)
        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Sale id'
        hdr_cells[1].text = 'Date and time'
        for i in range(len(result)):
            if len(result) == 0:
                document.add_paragraph("No sales yet")

            sale_id = re.search('ID: (.*), seller id', result[i]).group(1)
            sale_date = re.search(', date:(.*)', result[i]).group(1)
            row_cells = table.add_row().cells
            row_cells[0].text = sale_id
            row_cells[1].text = sale_date

        document.save("Reports\\" + result_id + "." + result_name + ".docx")

