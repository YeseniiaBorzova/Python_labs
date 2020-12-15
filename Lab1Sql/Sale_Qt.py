from PyQt5 import QtWidgets, uic
import sys
import re
from PyQt5.QtWidgets import QMessageBox
from docx import Document
import openpyxl
from PyQt5.QtGui import QStandardItem, QStandardItemModel
from PyQt5.QtWidgets import QAbstractItemView, QMainWindow, QApplication, QPushButton, QVBoxLayout, QLabel, QWidget

from Sales import Sales
from Seller import Seller
from Product import Product


class Sales_Window(QMainWindow):
    def __init__(self):
        super().__init__()
        uic.loadUi("sale_form.ui", self)
        self.searchSaleBtn.clicked.connect(self.find_sale)
        self.addSale.clicked.connect(self.add_sale)
        self.updSales.clicked.connect(self.update_all_sales)
        self.delSale.clicked.connect(self.delete_sale)
        self.saleDocx.clicked.connect(self.form_docx_report)
        self.saleExcel.clicked.connect(self.form_exel_report)
        self.update_all_sales()
        self.show_all_sellers()
        self.show_all_products()

    def update_all_sales(self):
        self.sale_list.setEditTriggers(QAbstractItemView.NoEditTriggers)
        model = QStandardItemModel(self.sale_list)
        for i in Sales.show_all_sales().split("\n"):
            item = QStandardItem(i)
            model.appendRow(item)
        self.sale_list.setModel(model)

    def delete_sale(self):
        try:
            id = int(self.idEnter.text())
            Sales.delete_sale_by_id(id)
            self.update_all_sales()
            self.idEnter.setText("")
        except:
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Critical)
            msg.setText("Error")
            msg.setInformativeText('Please insert sale id!')
            msg.setWindowTitle("Error")
            msg.exec_()

    def show_all_sellers(self):
        self.allSellers.setEditTriggers(QAbstractItemView.NoEditTriggers)
        model = QStandardItemModel(self.allSellers)
        for i in Seller.show_all_sellers().split("\n"):
            item = QStandardItem(i)
            model.appendRow(item)
        self.allSellers.setModel(model)

    def show_all_products(self):
        self.allProds.setEditTriggers(QAbstractItemView.NoEditTriggers)
        model = QStandardItemModel(self.allProds)
        for i in Product.show_all_the_products().split("\n"):
            item = QStandardItem(i)
            model.appendRow(item)
        self.allProds.setModel(model)

    def find_sale(self):
        try:
            model = QStandardItemModel(self.search_sale_list)
            item = QStandardItem(Sales.find_sale_by_date(self.choseDate.date().toString("dd.MM.yyyy")))
            model.appendRow(item)
            self.search_sale_list.setModel(model)
        except:
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Critical)
            msg.setText("Error")
            msg.setInformativeText('Please insert data!')
            msg.setWindowTitle("Error")
            msg.exec_()

    def add_sale(self):
        try:
            arr_prod = self.products.text().split(",")
            arr_amount = list(map(float, self.amounts.text().split(",")))
            new_sale = Sales(int(self.seller_id.text()), arr_prod, arr_amount)
            new_sale.init_sale()
            self.update_all_sales()
            self.seller_id.setText("")
            self.products.setText("")
            self.amounts.setText("")
        except:
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Critical)
            msg.setText("Error")
            msg.setInformativeText('Please insert data!')
            msg.setWindowTitle("Error")
            msg.exec_()

    def form_exel_report(self):
        wb = openpyxl.load_workbook('Reports\\All_info.xlsx')
        if 'Sales' not in wb.sheetnames:
            wb.create_sheet("Sales")
        ws = wb["Sales"]
        ws.delete_cols(1, 5)
        ws.delete_rows(1, 100)
        sales_list = Sales.show_all_sales().split("\n")
        del sales_list[-1]
        for i in range(len(sales_list)):
            result_id = re.search('Id: (.*), seller id:', sales_list[i]).group(1)
            result_seller_id = re.search('seller id: (.*), ', sales_list[i]).group(1)
            result_date = re.search('datetime:(.*)', sales_list[i]).group(1)
            ws.cell(row=i + 1, column=1).value = result_id
            ws.cell(row=i + 1, column=2).value = result_seller_id
            ws.cell(row=i + 1, column=3).value = result_date

        wb.save('Reports\\All_info.xlsx')

    def form_docx_report(self):
        sales_list = Sales.show_all_sales().split("\n")
        del sales_list[-1]
        for i in sales_list:
            self.sale_export(i)

    def sale_export(self, sale):
        result_id = re.search('Id: (.*), seller id:', sale).group(1)
        result_seller_id = re.search('seller id: (.*), ', sale).group(1)
        result_date = re.search('datetime:(.*)', sale).group(1)
        document = Document()
        document.add_heading((result_id + " (Sale)"), 0)
        document.add_heading("Overall information", level=1)
        document.add_paragraph("Sale id: " + result_id)
        document.add_paragraph("Seller id: " + result_seller_id)
        document.add_paragraph("Date and time: " + result_date)

        document.save("Reports\\" + "id" + result_id + ".docx")

