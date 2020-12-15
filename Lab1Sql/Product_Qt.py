from PyQt5 import QtWidgets, uic
import sys
import re
from docx import Document
import openpyxl
from PyQt5.QtWidgets import QMessageBox
from PyQt5.QtGui import QStandardItem, QStandardItemModel
from PyQt5.QtWidgets import QAbstractItemView, QMainWindow, QApplication, QPushButton, QVBoxLayout, QLabel, QWidget

from Product import Product


class Manipulate_Product(QMainWindow):
    def __init__(self, name):
        super().__init__()
        uic.loadUi("manipulate_prod.ui", self)
        self.prod_name = name

        self.changeSellBtn.clicked.connect(self.change_sell_price)
        self.changePurchBtn.clicked.connect(self.change_purchase_price)
        self.changeNameBtn.clicked.connect(self.change_name)
        self.delBtn.clicked.connect(self.delete_product)

    def delete_product(self):
        try:
            Product.delete_product_by_name(self.prod_name)
            self.close()
        except:
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Critical)
            msg.setText("Error")
            msg.setInformativeText('Please insert data!!!')
            msg.setWindowTitle("Error")
            msg.exec_()

    def change_sell_price(self):
        try:
            Product.update_sell_price(self.prod_name, float(self.enterSell.text()))
            self.enterSell.setText("")
        except:
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Critical)
            msg.setText("Error")
            msg.setInformativeText('Please insert data!!!')
            msg.setWindowTitle("Error")
            msg.exec_()

    def change_purchase_price(self):
        try:
            Product.update_purchase_price(self.prod_name, float(self.enterPurch.text()))
            self.enterPurch.setText("")
        except:
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Critical)
            msg.setText("Error")
            msg.setInformativeText('Please insert data!!!')
            msg.setWindowTitle("Error")
            msg.exec_()

    def change_name(self):
        try:
            Product.update_product_name(self.prod_name, self.enterName.text())
            self.enterName.setText("")
        except:
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Critical)
            msg.setText("Error")
            msg.setInformativeText('Please insert data!!!')
            msg.setWindowTitle("Error")
            msg.exec_()


class Product_Window(QMainWindow):
    def __init__(self):
        super().__init__()
        uic.loadUi("product_form.ui", self)
        self.searchProdBtn.clicked.connect(self.find_product)
        self.add_prod.clicked.connect(self.add_product)
        self.manipulate_prod.clicked.connect(self.manipulate_product)
        self.updBtn.clicked.connect(self.update_all_products)
        self.productDocx.clicked.connect(self.form_docx_report)
        self.productExcel.clicked.connect(self.form_exel_report)
        self.update_all_products()

    def update_all_products(self):
        self.all_products_list.setEditTriggers(QAbstractItemView.NoEditTriggers)
        model = QStandardItemModel(self.all_products_list)
        product_list = Product.show_all_the_products().split("\n")
        del product_list[-1]
        for i in product_list:
            item = QStandardItem(i)
            model.appendRow(item)
        self.all_products_list.setModel(model)

    def find_product(self):
        try:
            model = QStandardItemModel(self.search_product_list)
            item = QStandardItem(Product.find_product_by_name(self.enter_prod_name.text()))
            model.appendRow(item)
            self.search_product_list.setModel(model)
            self.enter_prod_name.setText("")
        except:
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Critical)
            msg.setText("Error")
            msg.setInformativeText('No such product')
            msg.setWindowTitle("Error")
            msg.exec_()

    def add_product(self):
        try:
            if self.prod_name.text() == "" or self.prod_unit.text() == "":
                msg = QMessageBox()
                msg.setIcon(QMessageBox.Critical)
                msg.setText("Error")
                msg.setInformativeText('Please insert data!')
                msg.setWindowTitle("Error")
                msg.exec_()
            else:
                new_prod = Product(self.prod_name.text(), self.prod_unit.text(),self.purchase_price.text(), self.sell_price.text())
                new_prod.insert_product()
                self.update_all_products()
                self.prod_name.setText("")
                self.prod_unit.setText("")
                self.purchase_price.setText("")
                self.sell_price.setText("")
        except:
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Critical)
            msg.setText("Error")
            msg.setInformativeText('Please insert data!')
            msg.setWindowTitle("Error")
            msg.exec_()

    def manipulate_product(self):
        name = self.manipulate_name.text()
        self.window = Manipulate_Product(name)
        self.window.show()
        self.update_all_products()
        self.manipulate_name.setText("")

    def form_docx_report(self):
        product_list = Product.show_all_the_products().split("\n")
        del product_list[-1]
        for i in product_list:
            self.product_export(i)

    def product_export(self, prod):
        result_id = re.search('Id: (.*), name', prod).group(1)
        result_name = re.search('name: (.*); ', prod).group(1)
        result_unit = re.search(' unit: (.*), purchase', prod).group(1)
        result_purch_price = re.search('purchase price: (.*), sell', prod).group(1)
        result_sell = re.search('sell price:(.*)', prod).group(1)

        document = Document()
        document.add_heading((result_id + ' ' + result_name + ' ' + "(Product)"), 0)
        document.add_heading("Overall information", level=1)
        document.add_paragraph("Id: " + result_id)
        document.add_paragraph("Name: " + result_name)
        document.add_paragraph("Measurement unit: " + result_unit)
        document.add_paragraph("Purchase price: " + result_purch_price)
        document.add_paragraph("Sell price: " + result_sell)

        document.save("Reports\\" + result_id + "." + result_name + ".docx")

    def form_exel_report(self):
        wb = openpyxl.load_workbook('Reports\\All_info.xlsx')
        if 'Products' not in wb.sheetnames:
            wb.create_sheet("Products")
        ws = wb["Products"]
        ws.delete_cols(1, 5)
        ws.delete_rows(1, 100)
        prod_list = Product.show_all_the_products().split("\n")
        del prod_list[-1]
        for i in range(len(prod_list)):
            result_id = re.search('Id: (.*), name', prod_list[i]).group(1)
            result_name = re.search('name: (.*); ', prod_list[i]).group(1)
            result_unit = re.search(' unit: (.*), purchase', prod_list[i]).group(1)
            result_purch_price = re.search('purchase price: (.*), sell', prod_list[i]).group(1)
            result_sell = re.search('sell price:(.*)', prod_list[i]).group(1)
            ws.cell(row=i + 1, column=1).value = result_id
            ws.cell(row=i + 1, column=2).value = result_name
            ws.cell(row=i + 1, column=3).value = result_unit
            ws.cell(row=i + 1, column=4).value = result_purch_price
            ws.cell(row=i + 1, column=5).value = result_sell

        wb.save('Reports\\All_info.xlsx')
