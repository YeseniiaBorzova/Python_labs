from Seller import Seller
from page import Page
from tkinter import *
from docx import Document
import openpyxl
from functools import partial
from tkinter import messagebox
import re


class SellerPage(Page):

    def __init__(self, *args, **kwargs):
        Page.__init__(self, *args, **kwargs)

        window = Toplevel(self)
        window.title("Seller window")
        window.state('zoomed')

        all_sellers = Frame(window)
        all_sellers.pack(side=LEFT, fill=Y)
        Label(all_sellers, text="All avaliable sellers").pack()
        self.allseller_listbox = Listbox(all_sellers, width=80)
        self.allseller_listbox.bind('<Double-1>', self.all_list_on_click)
        self.seller_list = Seller.show_all_sellers().split("\n")
        self.seller_list.pop(-1)
        for seller in self.seller_list:
            self.allseller_listbox.insert(END, seller)
        self.allseller_listbox.pack(side="top", fill="both")
        all_docx = Button(all_sellers, text="Export info in docx about all sellers", command=self.form_docx_report())
        all_docx.pack()
        all_exel = Button(all_sellers, text="Export info in exel about all sellers", command = self.form_exel_report())
        all_exel.pack()

        search_seller = Frame(window)
        search_seller.pack(side=LEFT, fill=Y)
        Label(search_seller, text="Search seller by surname").pack()
        enter_name = Entry(search_seller)
        enter_name.pack()
        self.find_listbox = Listbox(search_seller, width=80)
        self.find_listbox.bind('<Double-1>', self.find_list_on_click)
        find_prod = partial(self.search_seller_on_click, self.find_listbox, enter_name)
        search_btn = Button(search_seller, text="Search", command=find_prod)
        search_btn.pack()
        self.find_listbox.pack()

        add_seller = Frame(window)
        add_seller.pack(side=LEFT, fill=Y)
        Label(add_seller, text="Add new seller").grid(row=0, column=0, columnspan=3)
        name_label = Label(add_seller, text="Enter sellers` name:")
        surname_label = Label(add_seller, text="Enter sellers` surname:")
        middle_name_label = Label(add_seller, text="Enter sellers middle name:")
        income_percent_label = Label(add_seller, text="Enter sellers` income price:")
        name_entry = Entry(add_seller)
        surname_entry = Entry(add_seller)
        middle_name_entry = Entry(add_seller)
        income_percent_entry = Entry(add_seller)
        add_sel = partial(self.add_seller,  name_entry, surname_entry, middle_name_entry, income_percent_entry, self.allseller_listbox)
        add_btn = Button(add_seller, text="Add", command=add_sel)
        name_label.grid(row=1, column=0)
        surname_label.grid(row=2, column=0)
        middle_name_label.grid(row=3, column=0)
        income_percent_label.grid(row=4, column=0)
        name_entry.grid(row=1, column=1)
        surname_entry.grid(row=2, column=1)
        middle_name_entry.grid(row=3, column=1)
        income_percent_entry.grid(row=4, column=1)
        add_btn.grid(row=5, column=1)

    def all_list_on_click(self, event):
        cs = self.allseller_listbox.curselection()
        self.manipulate_seller(self.seller_list[cs[0]])

    def find_list_on_click(self, event):
        cs = self.find_listbox.curselection()
        self.manipulate_seller(self.find_list)

    def search_seller_on_click(self, listbox, entry_surname):
        listbox.delete(0, END)
        try:
            self.find_list = Seller.find_seller_by_surname(entry_surname.get()) + self.find_income(entry_surname)
            listbox.insert(END, self.find_list)
        except IndexError:
            messagebox.showerror("No such seller error!", "Incorrect seller surname or such seller doesnt exist!")
        return

    def update_list(self, listbox):
        listbox.delete(0, END)
        self.seller_list = Seller.show_all_sellers().split('\n')
        for seller in self.seller_list:
            self.allseller_listbox.insert(END, seller)
        return

    def manipulate_seller(self, seller_name):
        word_list = re.sub("[^\w]", " ", seller_name).split()
        chosen_name = word_list[3]
        chosen_surname = word_list[4]
        new_window = Toplevel(self)
        new_window.title("Seller manipulation window")
        new_window.geometry("800x200")

        change_percent = Frame(new_window)
        change_percent.pack(side=LEFT, fill=Y)
        Label(change_percent, text=chosen_surname+' '+chosen_name).grid(columnspan=2)
        Label(change_percent, text="Change employees commission percent here").grid(row=1,column=2)
        percent_label = Label(change_percent,text="Enter new percent:")
        percent_label.grid(row=2, column=1)
        percent_entry = Entry(change_percent)
        percent_entry.grid(row=2, column=2)
        change_percent_fun = partial(self.update_percent, chosen_surname,percent_entry)
        percent_btn = Button(change_percent,text="Change", command=change_percent_fun)
        percent_btn.grid(row=3, column=2)

    def update_percent(self, surname, percent):
        Seller.update_percent(surname, float(percent.get()))
        self.update_list(self.allseller_listbox)

    def add_seller(self,name_entry, surname_entry, middle_name_entry, income_percent_entry, listbox):
        new_seller = Seller(name_entry.get(), surname_entry.get(), middle_name_entry.get(), float(income_percent_entry.get()))
        new_seller.insert_seller()
        name_entry.delete(0, END)
        surname_entry.delete(0, END)
        middle_name_entry.delete(0, END)
        income_percent_entry.delete(0, END)
        self.update_list(listbox)

    def find_income(self, surname):
        self.seller_income = Seller.seller_income_by_surname(surname.get())
        return ", income: "+self.seller_income

    def form_exel_report(self):
        wb = openpyxl.load_workbook('Reports\\All_info.xlsx')
        if 'Sellers' not in wb.sheetnames:
            wb.create_sheet("Sellers")
        ws = wb.get_sheet_by_name("Sellers")
        ws.delete_cols(1, 3)
        ws.delete_rows(1, 100)
        for i in range(len(self.seller_list)):
            result_id = re.search('Id: (.*), Name', self.seller_list[i]).group(1)
            result_name = re.search('Name: (.*), comission ', self.seller_list[i]).group(1)
            result_percent = re.search(' percent: (.*) ', self.seller_list[i]).group(1)
            ws.cell(row=i + 1, column=1).value = result_id
            ws.cell(row=i + 1, column=2).value = result_name
            ws.cell(row=i + 1, column=3).value = result_percent


        wb.save('Reports\\All_info.xlsx')

    def form_docx_report(self):
        for i in self.seller_list:
            self.seller_export(i)

    def seller_export(self, seller):
        result_id = re.search('Id: (.*), Name', seller).group(1)
        result_name = re.search('Name: (.*), comission ', seller).group(1)
        result_percent = re.search(' percent: (.*) ', seller).group(1)
        result = Seller.find_seller_sales(result_id).split("\n")
        result.pop(-1)
        document = Document()
        document.add_heading((result_id + ' ' + result_name + ' ' + "(Seller)"), 0)
        document.add_heading("Overall information", level=1)
        document.add_paragraph("Id: " + result_id)
        document.add_paragraph("Name: " + result_name)
        document.add_paragraph("Commission percent: " + result_percent)
        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Sale id'
        hdr_cells[1].text = 'Date and time'
        for i in range(len(result)):
            if len(result)==0:
                document.add_paragraph("No sales yet")

            sale_id = re.search('ID: (.*), seller id', result[i]).group(1)
            sale_date = re.search(', date:(.*)', result[i]).group(1)
            row_cells = table.add_row().cells
            row_cells[0].text = sale_id
            row_cells[1].text = sale_date

        document.save("Reports\\" + result_id + "." + result_name + ".docx")