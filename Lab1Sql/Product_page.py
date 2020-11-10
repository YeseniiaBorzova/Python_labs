from functools import partial
import re
from docx import Document
import openpyxl
from Product import Product
from page import Page
from tkinter import messagebox
from tkinter import *


class ProductPage(Page):

    def __init__(self, *args, **kwargs):
        Page.__init__(self, *args, **kwargs)

        window = Toplevel(self)
        window.title("Product window")
        window.state('zoomed')

        all_products = Frame(window)
        all_products.pack(side=LEFT, fill=Y)
        Label(all_products, text="All avaliable products").pack()
        self.allprod_listbox = Listbox(all_products, width=80)
        self.allprod_listbox.bind('<Double-1>',self.all_list_on_click)
        self.prod_list = Product.show_all_the_products().split("\n")
        self.prod_list.pop(-1)
        for prod in self.prod_list:
            self.allprod_listbox.insert(END, prod)
        self.allprod_listbox.pack(side="top", fill="both")
        all_docx = Button(all_products, text="Export info in docx about all products", command=self.form_docx_report())
        all_docx.pack()
        all_exel = Button(all_products, text="Export info in exel about all products", command=self.form_exel_report())
        all_exel.pack()

        search_product = Frame(window)
        search_product.pack(side=LEFT, fill=Y)
        Label(search_product, text="Search product by name").pack()
        enter_name = Entry(search_product)
        enter_name.pack()
        self.find_listbox = Listbox(search_product, width=80)
        self.find_listbox.bind('<Double-1>', self.find_list_on_click)
        find_prod = partial(self.search_prod_on_click, self.find_listbox, enter_name)
        search_btn = Button(search_product, text="Search", command=find_prod)
        search_btn.pack()
        self.find_listbox.pack()

        add_product = Frame(window)
        add_product.pack(side=LEFT, fill=Y)
        Label(add_product, text="Add new product").grid(row=0, column=0, columnspan=3)
        name_label = Label(add_product, text="Enter products` name:")
        unit_label = Label(add_product, text="Enter products` measurment unit:")
        sell_label = Label(add_product, text="Enter products` sell price:")
        purchase_label = Label(add_product, text="Enter products` purchase price:")
        name_entry = Entry(add_product)
        unit_entry = Entry(add_product)
        sell_entry = Entry(add_product)
        purchase_entry = Entry(add_product)
        add_prod = partial(self.add_product, name_entry, unit_entry, sell_entry, purchase_entry,self.allprod_listbox)
        add_btn = Button(add_product, text="Add", command=add_prod)
        name_label.grid(row=1, column=0)
        unit_label.grid(row=2, column=0)
        sell_label.grid(row=3, column=0)
        purchase_label.grid(row=4, column=0)
        name_entry.grid(row=1, column=1)
        unit_entry.grid(row=2, column=1)
        sell_entry.grid(row=3, column=1)
        purchase_entry.grid(row=4, column=1)
        add_btn.grid(row=5, column=1)

    def update_list(self, listbox):
        listbox.delete(0, END)
        self.prod_list = Product.show_all_the_products().split("\n")
        for prod in self.prod_list:
            self.allprod_listbox.insert(END, prod)
        return

    def search_prod_on_click(self, listbox, entry_name):
        listbox.delete(0,END)
        try:
            self.find_list = Product.find_product_by_name(entry_name.get())
            listbox.insert(END, self.find_list)
        except IndexError:
            messagebox.showerror("No such product error!", "Incorrect product name or such product doesnt exist!")
        return

    def add_product(self,name, measurment_unit, purchase_price, sell_price, listbox):
        new_prod = Product(name.get(), measurment_unit.get(), float(purchase_price.get()), float(sell_price.get()))
        new_prod.insert_product()
        name.delete(0, END)
        measurment_unit.delete(0, END)
        purchase_price.delete(0, END)
        sell_price.delete(0, END)
        self.update_list(listbox)

    def find_list_on_click(self, event):
        cs = self.find_listbox.curselection()
        self.manipulate_product(self.find_list)

    def all_list_on_click(self, event):
        cs = self.allprod_listbox.curselection()
        self.manipulate_product(self.prod_list[cs[0]])

    def manipulate_product(self, product_name):
        word_list = re.sub("[^\w]", " ", product_name).split()
        chosen_name = word_list[3]
        new_window = Toplevel(self)
        new_window.title("Product manipulation window")
        new_window.geometry("800x200")
        change_price = Frame(new_window)
        change_price.pack(side=LEFT, fill=Y)
        Label(change_price, text=product_name).grid(columnspan=2)
        Label(change_price, text="Change products` price here").grid(row=1, column=1)
        sell_label = Label(change_price, text="Enter new products` sell price:")
        purchase_label = Label(change_price, text="Enter new products` purchase price:")
        sell_entry = Entry(change_price)
        purchase_entry = Entry(change_price)
        sell_label.grid(row=2, column=0)
        purchase_label.grid(row=3, column=0)
        sell_entry.grid(row=2, column=1)
        purchase_entry.grid(row=3, column=1)

        sell_price_change = partial(self.update_sell_price, chosen_name, sell_entry)
        purchase_price_change = partial(self.update_purchase_price, chosen_name, purchase_entry)
        change_sell_price = Button(change_price, text="Change", command=sell_price_change)
        change_purchase_price = Button(change_price, text="Change", command=purchase_price_change)
        change_sell_price.grid(row=2, column=2)
        change_purchase_price.grid(row=3, column=2)

        change_name = Frame(new_window)
        change_name.pack(side=LEFT, fill=Y)
        Label(change_name, text="Change products` name here").grid(row=1, column=1, columnspan=3)
        name_label = Label(change_name, text="Enter new name")
        new_name_entry = Entry(change_name)
        change_product_name = partial(self.update_prod_name, chosen_name, new_name_entry)
        change_name_btn = Button(change_name, text="Change", command=change_product_name)
        name_label.grid(row=2, column=1)
        new_name_entry.grid(row=2, column=2)
        change_name_btn.grid(row=2, column=3)

    def update_sell_price(self, name, price):
        Product.update_sell_price(name, float(price.get()))
        self.update_list(self.allprod_listbox)

    def update_purchase_price(self, name, price):
        Product.update_purchase_price(name, float(price.get()))
        self.update_list(self.allprod_listbox)

    def update_prod_name(self, old_name, new_name):
        Product.update_product_name(old_name, new_name.get())
        self.update_list(self.allprod_listbox)

    def delete_prod(self, name):
        Product.delete_product_by_name(name.get)
        self.update_list(self.allprod_listbox)

    def form_exel_report(self):
        wb = openpyxl.load_workbook('Reports\\All_info.xlsx')
        if 'Products' not in wb.sheetnames:
            wb.create_sheet("Products")
        ws = wb.get_sheet_by_name("Products")
        ws.delete_cols(1,5)
        ws.delete_rows(1,100)
        for i in range(len(self.prod_list)):
            result_id = re.search('Id: (.*), name', self.prod_list[i]).group(1)
            result_name = re.search('name: (.*); ', self.prod_list[i]).group(1)
            result_unit = re.search(' unit: (.*), purchase', self.prod_list[i]).group(1)
            result_purch_price = re.search('purchase price: (.*), sell', self.prod_list[i]).group(1)
            result_sell = re.search('sell price:(.*)', self.prod_list[i]).group(1)
            ws.cell(row=i + 1, column=1).value = result_id
            ws.cell(row=i + 1, column=2).value = result_name
            ws.cell(row=i + 1, column=3).value = result_unit
            ws.cell(row=i + 1, column=4).value = result_purch_price
            ws.cell(row=i + 1, column=5).value = result_sell

        wb.save('Reports\\All_info.xlsx')

    def form_docx_report(self):
        for i in self.prod_list:
            self.prod_export(i)

    def prod_export(self, prod):
        result_id = re.search('Id: (.*), name', prod).group(1)
        result_name = re.search('name: (.*); ', prod).group(1)
        result_unit = re.search(' unit: (.*), purchase', prod).group(1)
        result_purch_price = re.search('purchase price: (.*), sell', prod).group(1)
        result_sell = re.search('sell price:(.*)', prod).group(1)

        document = Document()
        document.add_heading((result_id +' '+ result_name + ' ' + "(Product)"), 0)
        document.add_heading("Overall information", level=1)
        document.add_paragraph("Id: "+result_id)
        document.add_paragraph("Name: " + result_name)
        document.add_paragraph("Measurement unit: " + result_unit)
        document.add_paragraph("Purchase price: " + result_purch_price)
        document.add_paragraph("Sell price: " + result_sell)

        document.save("Reports\\" + result_id + "." + result_name + ".docx")
